import os
from pathlib import Path
import re

from docx import Document
from flask import Flask, render_template, request, send_file

from storage import engagement_artifacts

app = Flask(__name__)
BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE_DIR / "templates"
OUTPUT_DIR = BASE_DIR / "output"
OUTPUT_DIR.mkdir(exist_ok=True)


@app.route("/")
def form() -> str:
    return render_template("form.html")


@app.route("/generate", methods=["POST"])
def generate_document():
    form_data = request.form

    client_name = form_data.get("client_name", "").strip()
    client_type = form_data.get("client_type", "").strip()
    matter_type = form_data.get("matter_type", "").strip()
    payment_method = form_data.get("payment_method", "").strip()
    matter_description = form_data.get("matter_description", "").strip()
    instructing_officer_name = form_data.get("instructing_officer_name", "").strip()
    retainer_amount = form_data.get("retainer_amount", "").strip()
    client_type_norm = (client_type or "").strip().lower()
    matter_type_norm = (matter_type or "").strip().lower()
    payment_method_norm = (payment_method or "").strip().lower()

    if not client_name:
        return "Client name is required.", 400
    if not client_type:
        return "Client type is required.", 400
    if not matter_type:
        return "Matter type is required.", 400
    if not payment_method:
        return "Payment method is required.", 400
    if not matter_description:
        return "Matter description is required.", 400

    if client_type == "Corporation" and not instructing_officer_name:
        return "Instructing officer name is required for corporations.", 400
    if matter_type == "Hourly Strategy" and not retainer_amount:
        return "Retainer amount is required for Hourly Strategy matters.", 400

    if payment_method_norm == "retainer":
        if client_type_norm == "individual":
            template_filename = "RETAINERAgreement_Individual_v2.docx"
        elif client_type_norm == "corporation":
            template_filename = "RETAINERAgreement_Corporation_v2.docx"
        else:
            template_filename = None
    else:
        if client_type_norm == "individual" and matter_type_norm == "flat":
            template_filename = "ENGAGEMENTAgreement_Individual_[Flat]_v2.docx"
        elif client_type_norm == "individual" and matter_type_norm in {
            "hourly solution",
            "hourly strategy",
            "hourly",
        }:
            template_filename = "ENGAGEMENTAgreement_Individual_[Hourly]_v2.docx"
        elif client_type_norm == "corporation":
            template_filename = "ENGAGEMENTAgreement_Corporation_v2.docx"
        else:
            template_filename = None

    templates_dir = os.fspath(TEMPLATES_DIR)
    available_templates = sorted(
        name
        for name in os.listdir(templates_dir)
        if name.lower().endswith(".docx")
    )
    received_values = {
        "client_name": client_name,
        "client_type": client_type,
        "matter_type": matter_type,
        "payment_method": payment_method,
        "matter_description": matter_description,
        "instructing_officer_name": instructing_officer_name,
        "retainer_amount": retainer_amount,
    }

    if not template_filename:
        message = (
            "Unsupported selection: "
            f"client_type={client_type_norm}, "
            f"matter_type={matter_type_norm}, "
            f"payment_method={payment_method_norm}\n"
            f"available_templates: {available_templates}\n"
            f"received_values: {received_values}\n"
        )
        return message, 400

    template_path = os.path.join(templates_dir, template_filename)
    if not os.path.exists(template_path):
        message = (
            "Template not found.\n"
            f"template_path: {template_path}\n"
            f"available_templates: {available_templates}\n"
            f"received_values: {received_values}\n"
        )
        return message, 500

    replacements = {
        "client_name": client_name or "N/A",
        "client_type": client_type or "N/A",
        "instructing_officer_name": instructing_officer_name or "N/A",
        "matter_type": matter_type or "N/A",
        "retainer_amount": retainer_amount or "N/A",
        "payment_method": payment_method or "N/A",
        "matter_description": matter_description or "N/A",
    }

    document = Document(template_path)
    replace_placeholders(document, replacements)

    filename = f"{sanitize_filename(client_name)}.docx"
    file_path = OUTPUT_DIR / filename
    document.save(file_path)

    return send_file(file_path, as_attachment=True, download_name=filename)


def sanitize_filename(value: str) -> str:
    sanitized = re.sub(r"[^A-Za-z0-9_-]+", "_", value.strip())
    sanitized = sanitized.strip("_")
    return sanitized or "document"


def replace_placeholders(document: Document, replacements: dict) -> None:
    for paragraph in iter_paragraphs(document):
        replace_in_paragraph(paragraph, replacements)


def replace_in_paragraph(paragraph, replacements: dict) -> None:
    if not paragraph.runs:
        return
    original_text = "".join(run.text for run in paragraph.runs)
    updated_text = original_text
    for key, value in replacements.items():
        updated_text = updated_text.replace(f"{{{{{key}}}}}", value)
    if updated_text != original_text:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = updated_text


def iter_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def get_engagement_artifacts(engagement_id: str) -> dict:
    return engagement_artifacts(engagement_id, version=1)


if __name__ == "__main__":
    app.run(debug=True)
