import tempfile
import unittest
from pathlib import Path

from docx import Document

import app as app_module


class AppTestCase(unittest.TestCase):
    def setUp(self):
        self.app = app_module.app
        self.app.config["TESTING"] = True
        self.client = self.app.test_client()
        self.temp_dir = tempfile.TemporaryDirectory()
        app_module.OUTPUT_DIR = Path(self.temp_dir.name)
        app_module.OUTPUT_DIR.mkdir(exist_ok=True)
        self.template_dir = Path(self.temp_dir.name) / "templates"
        self.template_dir.mkdir(exist_ok=True)
        app_module.TEMPLATES_DIR = self.template_dir
        self.template_names = [
            "ENGAGEMENTAgreement_Individual_[Flat]_v2.docx",
            "ENGAGEMENTAgreement_Individual_[Hourly]_v2.docx",
            "ENGAGEMENTAgreement_Corporation_v2.docx",
            "RETAINERAgreement_Individual_v2.docx",
            "RETAINERAgreement_Corporation_v2.docx",
        ]
        for template_name in self.template_names:
            self._create_template(self.template_dir / template_name, template_name)

    def tearDown(self):
        self.temp_dir.cleanup()

    def _create_template(self, path: Path, template_label: str) -> None:
        document = Document()
        document.add_paragraph(f"Template: {template_label}")
        document.add_paragraph("Client: {{client_name}}")
        document.add_paragraph("Officer: {{instructing_officer_name}}")
        document.add_paragraph("Retainer: {{retainer_amount}}")
        document.save(path)

    def test_get_form_contains_fields(self):
        response = self.client.get("/")
        self.assertEqual(response.status_code, 200)
        body = response.get_data(as_text=True)
        for field_name in [
            "client_name",
            "client_type",
            "matter_type",
            "payment_method",
            "matter_description",
            "instructing_officer_name",
            "retainer_amount",
        ]:
            self.assertIn(field_name, body)

    def test_requires_instructing_officer_for_corporation(self):
        response = self.client.post(
            "/generate",
            data={
                "client_name": "Acme",
                "client_type": "Corporation",
                "matter_type": "Flat",
                "payment_method": "retainer",
                "matter_description": "incorporation",
            },
        )
        self.assertEqual(response.status_code, 400)
        self.assertIn("Instructing officer name is required", response.get_data(as_text=True))

    def test_requires_retainer_amount_for_hourly_strategy(self):
        response = self.client.post(
            "/generate",
            data={
                "client_name": "Pat",
                "client_type": "Individual",
                "matter_type": "Hourly Strategy",
                "payment_method": "pay on invoice",
                "matter_description": "corporate advisory",
            },
        )
        self.assertEqual(response.status_code, 400)
        self.assertIn("Retainer amount is required", response.get_data(as_text=True))

    def test_valid_submission_generates_document(self):
        response = self.client.post(
            "/generate",
            data={
                "client_name": "Jordan",
                "client_type": "Individual",
                "matter_type": "Hourly Solution",
                "payment_method": "pay on invoice",
                "matter_description": "business acquisition",
            },
        )
        self.assertEqual(response.status_code, 200)
        response.close()
        output_files = list(app_module.OUTPUT_DIR.glob("Jordan.docx"))
        self.assertEqual(len(output_files), 1)

    def _assert_template_used(self, response, expected_template: str) -> None:
        self.assertEqual(response.status_code, 200)
        response.close()
        output_files = list(app_module.OUTPUT_DIR.glob("*.docx"))
        self.assertEqual(len(output_files), 1)
        document = Document(output_files[0])
        full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn(f"Template: {expected_template}", full_text)

    def test_template_mapping_individual_flat_invoice(self):
        response = self.client.post(
            "/generate",
            data={
                "client_name": "Alex",
                "client_type": "Individual",
                "matter_type": "Flat",
                "payment_method": "pay on invoice",
                "matter_description": "incorporation",
            },
        )
        self._assert_template_used(
            response, "ENGAGEMENTAgreement_Individual_[Flat]_v2.docx"
        )

    def test_template_mapping_individual_hourly_invoice(self):
        response = self.client.post(
            "/generate",
            data={
                "client_name": "Bailey",
                "client_type": "Individual",
                "matter_type": "Hourly Solution",
                "payment_method": "pay on invoice",
                "matter_description": "shareholder agreement",
            },
        )
        self._assert_template_used(
            response, "ENGAGEMENTAgreement_Individual_[Hourly]_v2.docx"
        )

    def test_template_mapping_individual_hourly_retainer(self):
        response = self.client.post(
            "/generate",
            data={
                "client_name": "Casey",
                "client_type": "Individual",
                "matter_type": "Hourly Strategy",
                "payment_method": "retainer",
                "matter_description": "business advisory",
                "retainer_amount": "5000",
            },
        )
        self._assert_template_used(response, "RETAINERAgreement_Individual_v2.docx")

    def test_template_mapping_corporation_flat_invoice(self):
        response = self.client.post(
            "/generate",
            data={
                "client_name": "Delta Co",
                "client_type": "Corporation",
                "matter_type": "Flat",
                "payment_method": "pay on invoice",
                "matter_description": "incorporation",
                "instructing_officer_name": "Taylor",
            },
        )
        self._assert_template_used(response, "ENGAGEMENTAgreement_Corporation_v2.docx")

    def test_template_mapping_corporation_hourly_retainer(self):
        response = self.client.post(
            "/generate",
            data={
                "client_name": "Echo Corp",
                "client_type": "Corporation",
                "matter_type": "Hourly Strategy",
                "payment_method": "retainer",
                "matter_description": "shareholder agreement",
                "instructing_officer_name": "Robin",
                "retainer_amount": "7500",
            },
        )
        self._assert_template_used(response, "RETAINERAgreement_Corporation_v2.docx")


if __name__ == "__main__":
    unittest.main()
